"""Workflow-local proposal DOCX builder with bid-document typography."""

from __future__ import annotations

import json
import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import Any


IMAGE_MARKER = '[[WORKFLOW_IMAGES]]'
HEADING = re.compile(r'^(#{1,6})\s+(.+?)\s*$')
BULLET = re.compile(r'^\s*[-*+]\s+(.+)$')
NUMBERED = re.compile(r'^\s*\d+[.)、]\s+(.+)$')
INLINE_IMAGE = re.compile(r'^!\[([^]]*)\]\(([^)]+)\)\s*$')


def _clean_inline(text: str) -> str:
    value = re.sub(r'!\[([^]]*)\]\([^)]+\)', r'\1', str(text or ''))
    value = re.sub(r'\[([^]]+)\]\([^)]+\)', r'\1', value)
    value = re.sub(r'(\*\*|__|~~|`)', '', value)
    return value.strip()


def _strip_heading_number(text: str) -> str:
    return re.sub(r'^\s*\d+(?:\.\d+){0,3}[、.．\s　]+', '', _clean_inline(text)).strip()


def _safe_output_name(value: str) -> str:
    name = Path(str(value or '')).name
    name = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', name).strip(' ._') or '投标技术方案.docx'
    return name if name.lower().endswith('.docx') else name + '.docx'


def _json_object(value: Any, name: str) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    try:
        parsed = json.loads(str(value))
    except (TypeError, json.JSONDecodeError) as exc:
        raise ValueError(f'{name} must be a valid JSON object.') from exc
    if not isinstance(parsed, dict):
        raise ValueError(f'{name} must be a JSON object.')
    return parsed


def _image_paths(value: Any) -> list[dict[str, str]]:
    if isinstance(value, list):
        items = value
    else:
        try:
            items = json.loads(str(value))
        except json.JSONDecodeError as exc:
            raise ValueError(f'image_paths_json is invalid: {exc.msg}') from exc
    if not isinstance(items, list):
        raise ValueError('image_paths_json must be a JSON list.')
    output: list[dict[str, str]] = []
    for index, item in enumerate(items[:12]):
        if isinstance(item, dict):
            path = str(item.get('path') or item.get('url') or '').strip()
            title = str(item.get('title') or item.get('caption') or f'方案图 {index + 1}').strip()
            kind = str(item.get('type') or item.get('image_type') or '').strip()
        else:
            path, title, kind = str(item).strip(), f'方案图 {index + 1}', ''
        resolved = Path(path).expanduser().resolve()
        if not resolved.is_file():
            raise FileNotFoundError(f'Image does not exist: {path}')
        output.append({'path': str(resolved), 'title': title, 'type': kind})
    if len(output) < 1:
        raise ValueError('At least one real proposal image is required.')
    return output


def _set_run_font(run: Any, east_asia: str = '宋体', latin: str = 'Times New Roman',
                  size: float | None = None, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _multilevel_numbering(document: Any) -> int:
    """Create a 1 / 1.1 / 1.1.1 / 1.1.1.1 heading numbering definition."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn('w:abstractNumId'))) for node in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(node.get(qn('w:numId'))) for node in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'multilevel')
    abstract.append(multi)
    formats = ('%1', '%1.%2', '%1.%2.%3', '%1.%2.%3.%4')
    sizes = (36, 32, 32, 28)  # half points: 18, 16, 16, 14 pt
    for level, (pattern, size) in enumerate(zip(formats, sizes)):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'decimal')
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), pattern)
        suffix = OxmlElement('w:suff')
        suffix.set(qn('w:val'), 'space')
        p_style = OxmlElement('w:pStyle')
        p_style.set(qn('w:val'), f'Heading{level + 1}')
        p_pr = OxmlElement('w:pPr')
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'num')
        tab.set(qn('w:pos'), str(360 + level * 360))
        tabs.append(tab)
        indent = OxmlElement('w:ind')
        indent.set(qn('w:left'), str(360 + level * 360))
        indent.set(qn('w:hanging'), '360')
        p_pr.extend([tabs, indent])
        r_pr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Arial')
        fonts.set(qn('w:eastAsia'), '微软雅黑')
        bold = OxmlElement('w:b')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size))
        r_pr.extend([fonts, bold, sz])
        lvl.extend([start, num_fmt, lvl_text, suffix, p_style, p_pr, r_pr])
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _configure_document(document: Any) -> int:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    for section in document.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Pt(28)
    normal.paragraph_format.space_after = Pt(0)
    sizes = {1: 18, 2: 16, 3: 16, 4: 14}
    for level, size in sizes.items():
        style = document.styles[f'Heading {level}']
        style.font.name = 'Arial'
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.page_break_before = level == 1
    if 'Bid Caption' not in [style.name for style in document.styles]:
        caption = document.styles.add_style('Bid Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption.base_style = normal
        caption.font.size = Pt(10.5)
        caption.paragraph_format.first_line_indent = Pt(0)
    num_id = _multilevel_numbering(document)
    for level in range(1, 5):
        style = document.styles[f'Heading {level}']
        p_pr = style.element.get_or_add_pPr()
        num_pr = p_pr.find(qn('w:numPr'))
        if num_pr is None:
            num_pr = OxmlElement('w:numPr')
            p_pr.append(num_pr)
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level - 1))
        number = OxmlElement('w:numId')
        number.set(qn('w:val'), str(num_id))
        num_pr.extend([ilvl, number])
    return num_id


def _add_toc(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(heading.add_run('目录'), east_asia='微软雅黑', latin='Arial', size=18, bold=True)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = 0
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = 'TOC \\o "1-4" \\h \\z \\u'
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = '请在 Word/WPS 中更新目录域'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)
    document.add_page_break()


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
    return bool(cells) and all(re.fullmatch(r':?-{3,}:?', cell or '') for cell in cells)


def _add_table(document: Any, lines: list[str]) -> None:
    rows = [[_clean_inline(cell.strip()) for cell in line.strip().strip('|').split('|')]
            for line in lines if not _is_table_separator(line)]
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = 'Table Grid'
    for row_index, row in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.text = ''
            text = row[column_index] if column_index < len(row) else ''
            run = cell.paragraphs[0].add_run(text)
            cell.paragraphs[0].paragraph_format.first_line_indent = 0
            _set_run_font(run, size=10.5, bold=row_index == 0)


def _add_images(document: Any, images: list[dict[str, str]]) -> int:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    count = 0
    for item in images:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = 0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            paragraph.add_run().add_picture(item['path'], width=Cm(15))
        except Exception as exc:
            raise ValueError(f"Cannot embed image {item['path']}: {exc}") from exc
        # The source skill's latest image guide explicitly avoids auto-generated
        # standalone figure captions. Titles remain in the image manifest.
        count += 1
    return count


def _render_markdown(document: Any, markdown_text: str, images: list[dict[str, str]]) -> tuple[int, int, int]:
    lines = str(markdown_text or '').replace('\r\n', '\n').replace('\r', '\n').split('\n')
    index = 0
    paragraph_count = 0
    heading_count = 0
    embedded = 0
    image_marker_seen = False
    in_code = False
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if stripped.startswith('```'):
            in_code = not in_code
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped == IMAGE_MARKER:
            embedded += _add_images(document, images)
            image_marker_seen = True
            index += 1
            continue
        if '|' in stripped and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip() and '|' in lines[index]:
                table_lines.append(lines[index])
                index += 1
            _add_table(document, table_lines)
            paragraph_count += len(table_lines) - 1
            continue
        match = HEADING.match(stripped)
        if match:
            level = min(4, len(match.group(1)))
            title = _strip_heading_number(match.group(2))
            document.add_heading(title, level=level)
            heading_count += 1
            index += 1
            continue
        if INLINE_IMAGE.match(stripped):
            index += 1
            continue
        bullet = BULLET.match(line)
        numbered = NUMBERED.match(line)
        if bullet or numbered:
            paragraph = document.add_paragraph(style='List Bullet' if bullet else 'List Number')
            paragraph.paragraph_format.first_line_indent = 0
            run = paragraph.add_run(_clean_inline((bullet or numbered).group(1)))
        else:
            paragraph = document.add_paragraph(style='Normal')
            run = paragraph.add_run(_clean_inline(stripped))
        if in_code:
            _set_run_font(run, east_asia='等线', latin='Courier New', size=10.5)
        paragraph_count += 1
        index += 1
    if images and not image_marker_seen:
        document.add_heading('系统架构与效果', level=1)
        embedded += _add_images(document, images)
    return paragraph_count, heading_count, embedded


def compose_proposal_docx(markdown_text: str, outline_json: str,
                          image_paths_json: str = '[]', output_name: str = '投标技术方案.docx',
                          add_toc: bool = True) -> dict[str, Any]:
    """Compose complete proposal Markdown and local images into a styled DOCX.

    Args:
        markdown_text: Complete proposal Markdown; [[WORKFLOW_IMAGES]] controls placement.
        outline_json: Validated outline object or JSON string used for title metadata.
        image_paths_json: Ordered JSON list of local path/title/type objects.
        output_name: Safe project-specific DOCX filename without a directory.
        add_toc: Insert an updateable Heading 1-4 Word/WPS table of contents.

    Returns:
        Metadata including the absolute path of the generated DOCX.
    """
    if not isinstance(markdown_text, str) or len(markdown_text.strip()) < 200:
        raise ValueError('markdown_text must contain the complete proposal.')
    outline = _json_object(outline_json, 'outline_json')
    images = _image_paths(image_paths_json)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError('python-docx is required for proposal composition.') from exc

    root = Path(tempfile.gettempdir()) / 'bid_tech_proposal_writer' / 'documents' / uuid.uuid4().hex
    root.mkdir(parents=True, exist_ok=True)
    output = root / _safe_output_name(output_name)
    document = Document()
    _configure_document(document)
    title = str(outline.get('project_full_name') or outline.get('project_name') or '').strip()
    if not title:
        first = next((line for line in markdown_text.splitlines() if line.strip()), '投标技术方案')
        title = _strip_heading_number(first.lstrip('# '))
    if not title.endswith('技术方案'):
        title += '技术方案'
    document.core_properties.title = title
    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.first_line_indent = 0
    cover.paragraph_format.space_before = Pt(150)
    _set_run_font(cover.add_run(title), east_asia='微软雅黑', latin='Arial', size=26, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = 0
    subtitle.paragraph_format.space_before = Pt(60)
    _set_run_font(subtitle.add_run('投标文件技术部分'), east_asia='宋体', size=16, bold=True)
    document.add_page_break()
    if add_toc:
        _add_toc(document)

    lines = markdown_text.splitlines()
    first_index = next((idx for idx, line in enumerate(lines) if line.strip()), None)
    if first_index is not None:
        first_match = HEADING.match(lines[first_index].strip())
        if first_match and _strip_heading_number(first_match.group(2)).rstrip('技术方案') == title.rstrip('技术方案'):
            lines[first_index] = ''
            # Shared Writer reserves Markdown H1 for the document title. The cover
            # already carries that title, so promote H2..H5 bid chapters to Word
            # Heading 1..4 after removing the duplicate Markdown title.
            for index, line in enumerate(lines):
                heading = HEADING.match(line.strip())
                if heading and len(heading.group(1)) >= 2:
                    lines[index] = line.replace(heading.group(1), heading.group(1)[1:], 1)
    paragraph_count, heading_count, embedded = _render_markdown(document, '\n'.join(lines), images)
    settings = document.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)
    document.save(output)
    if output.stat().st_size < 8000:
        raise RuntimeError(f'Generated DOCX is unexpectedly small: {output}')
    return {
        'path': str(output.resolve()), 'filename': output.name,
        'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'paragraph_count': paragraph_count, 'heading_count': heading_count,
        'images_embedded': embedded, 'toc': bool(add_toc),
        'renderer': 'workflow_local_python_docx',
    }
